import os
from flask import Flask, request, render_template, jsonify, send_file, session
import anthropic
from werkzeug.utils import secure_filename
import pdfplumber
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import traceback
from datetime import timedelta

app = Flask(__name__)

# Production configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DOWNLOAD_FOLDER'] = 'downloads'

# Use environment variable for secret key in production
app.secret_key = os.environ.get('SECRET_KEY', 'your-secret-key-change-this-in-production')
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=2)

# Ensure upload and download directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DOWNLOAD_FOLDER'], exist_ok=True)

# Initialize Anthropic client (will be initialized when needed)
client = None

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file_path, filename):
    """Extract text from uploaded resume file"""
    try:
        if filename.endswith('.pdf'):
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        
        elif filename.endswith('.docx'):
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        
        elif filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        
        else:
            return "Unsupported file format"
            
    except Exception as e:
        print(f"Error extracting text from {filename}: {str(e)}")
        traceback.print_exc()
        return f"Error reading file: {str(e)}"

def create_docx_resume(resume_content, filename_prefix="optimized_resume"):
    """Create a DOCX file from resume content with enhanced error handling"""
    try:
        if not resume_content or not resume_content.strip():
            raise ValueError("Resume content is empty or None")
        
        print(f"Creating resume document with {len(resume_content)} characters")
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Split content into lines and process
        lines = resume_content.strip().split('\n')
        print(f"Processing {len(lines)} lines for resume document")
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                # Add empty paragraph for spacing
                doc.add_paragraph()
                continue
            
            try:
                # Check if it's a header (all caps or starts with common header words)
                if (line.isupper() and len(line) > 3) or any(line.upper().startswith(header) for header in 
                    ['PROFESSIONAL', 'EXPERIENCE', 'EDUCATION', 'SKILLS', 'SUMMARY', 'OBJECTIVE', 'CONTACT', 'WORK', 'EMPLOYMENT']):
                    # Add as heading
                    heading = doc.add_heading(line, level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    # Add as bullet point
                    doc.add_paragraph(line[1:].strip(), style='List Bullet')
                else:
                    # Add as normal paragraph
                    doc.add_paragraph(line)
            except Exception as line_error:
                print(f"Error processing line {i}: {line_error}")
                # Fallback: add as plain paragraph
                doc.add_paragraph(line)
        
        # Save to temporary file
        file_id = str(uuid.uuid4())
        filename = f"{filename_prefix}_{file_id}.docx"
        filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
        
        print(f"Saving resume document to: {filepath}")
        doc.save(filepath)
        
        # Verify file was created and has content
        if not os.path.exists(filepath):
            raise Exception(f"Resume document was not saved to {filepath}")
        
        file_size = os.path.getsize(filepath)
        if file_size == 0:
            raise Exception(f"Resume document file is empty: {filepath}")
        
        print(f"Resume document created successfully: {filename} ({file_size} bytes)")
        return filename
        
    except Exception as e:
        print(f"Error creating resume document: {str(e)}")
        traceback.print_exc()
        raise e

def create_docx_cover_letter(cover_letter_content, filename_prefix="cover_letter"):
    """Create a DOCX file from cover letter content with enhanced error handling"""
    try:
        if not cover_letter_content or not cover_letter_content.strip():
            raise ValueError("Cover letter content is empty or None")
        
        print(f"Creating cover letter document with {len(cover_letter_content)} characters")
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run("Date: [Insert Date]")
        
        doc.add_paragraph()  # Empty line
        
        # Add placeholder for recipient info
        doc.add_paragraph("[Hiring Manager Name]")
        doc.add_paragraph("[Company Name]")
        doc.add_paragraph("[Company Address]")
        
        doc.add_paragraph()  # Empty line
        
        # Add subject line
        subject_para = doc.add_paragraph()
        subject_run = subject_para.add_run("Re: [Position Title] Application")
        subject_run.bold = True
        
        doc.add_paragraph()  # Empty line
        
        # Add cover letter content
        paragraphs = cover_letter_content.strip().split('\n\n')
        
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
                doc.add_paragraph()  # Add space between paragraphs
        
        # Add closing
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        doc.add_paragraph("[Your Name]")
        
        # Save to temporary file
        file_id = str(uuid.uuid4())
        filename = f"{filename_prefix}_{file_id}.docx"
        filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
        
        print(f"Saving cover letter document to: {filepath}")
        doc.save(filepath)
        
        # Verify file was created and has content
        if not os.path.exists(filepath):
            raise Exception(f"Cover letter document was not saved to {filepath}")
        
        file_size = os.path.getsize(filepath)
        if file_size == 0:
            raise Exception(f"Cover letter document file is empty: {filepath}")
        
        print(f"Cover letter document created successfully: {filename} ({file_size} bytes)")
        return filename
        
    except Exception as e:
        print(f"Error creating cover letter document: {str(e)}")
        traceback.print_exc()
        raise e

def get_anthropic_client():
    global client
    if client is None:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            print("ERROR: ANTHROPIC_API_KEY environment variable not found")
            print(f"Available environment variables: {list(os.environ.keys())}")
            raise ValueError("ANTHROPIC_API_KEY environment variable is required")
        
        print(f"API key found: {api_key[:10]}...{api_key[-4:] if len(api_key) > 14 else '***'}")
        try:
            client = anthropic.Anthropic(api_key=api_key)
            print("Anthropic client initialized successfully")
        except Exception as e:
            print(f"Failed to initialize Anthropic client: {e}")
            raise e
    return client

def parse_ai_response(result):
    """Enhanced parsing of AI response with multiple fallback methods"""
    print("Starting AI response parsing...")
    print(f"Response length: {len(result)} characters")
    
    resume_content = ''
    cover_letter_content = ''
    
    # Method 1: Look for ## markers
    if '## OPTIMIZED RESUME' in result and '## COVER LETTER' in result:
        print("Using ## marker parsing method")
        try:
            parts = result.split('## OPTIMIZED RESUME')[1]
            if '## COVER LETTER' in parts:
                resume_parts = parts.split('## COVER LETTER')
                resume_content = resume_parts[0].strip()
                cover_letter_content = resume_parts[1].strip()
                print(f"Method 1 success - Resume: {len(resume_content)}, Cover Letter: {len(cover_letter_content)}")
        except Exception as e:
            print(f"Method 1 failed: {e}")
    
    # Method 2: Look for single # markers if ## failed
    if (not resume_content or not cover_letter_content) and '# OPTIMIZED RESUME' in result and '# COVER LETTER' in result:
        print("Using # marker parsing method")
        try:
            parts = result.split('# OPTIMIZED RESUME')[1]
            if '# COVER LETTER' in parts:
                resume_parts = parts.split('# COVER LETTER')
                resume_content = resume_parts[0].strip()
                cover_letter_content = resume_parts[1].strip()
                print(f"Method 2 success - Resume: {len(resume_content)}, Cover Letter: {len(cover_letter_content)}")
        except Exception as e:
            print(f"Method 2 failed: {e}")
    
    # Method 3: Line-by-line parsing
    if not resume_content or not cover_letter_content:
        print("Using line-by-line parsing method")
        try:
            lines = result.split('\n')
            current_section = None
            resume_lines = []
            cover_letter_lines = []
            
            for line in lines:
                line_upper = line.upper().strip()
                if any(keyword in line_upper for keyword in ['OPTIMIZED RESUME', 'RESUME:', 'RESUME']):
                    current_section = 'resume'
                    print(f"Found resume section marker: {line}")
                    continue
                elif any(keyword in line_upper for keyword in ['COVER LETTER', 'LETTER:']):
                    current_section = 'cover_letter'
                    print(f"Found cover letter section marker: {line}")
                    continue
                
                if current_section == 'resume' and line.strip():
                    resume_lines.append(line)
                elif current_section == 'cover_letter' and line.strip():
                    cover_letter_lines.append(line)
            
            if not resume_content and resume_lines:
                resume_content = '\n'.join(resume_lines).strip()
            if not cover_letter_content and cover_letter_lines:
                cover_letter_content = '\n'.join(cover_letter_lines).strip()
            
            print(f"Method 3 result - Resume: {len(resume_content)}, Cover Letter: {len(cover_letter_content)}")
        except Exception as e:
            print(f"Method 3 failed: {e}")
    
    # Method 4: Split by keywords in the middle
    if not resume_content or not cover_letter_content:
        print("Using keyword split parsing method")
        try:
            # Look for patterns like "COVER LETTER" or "Cover Letter:" in the middle
            lower_result = result.lower()
            cover_letter_start = -1
            
            for pattern in ['cover letter:', 'cover letter', '2. cover letter', 'letter:']:
                pos = lower_result.find(pattern)
                if pos > 100:  # Make sure it's not at the very beginning
                    cover_letter_start = pos
                    break
            
            if cover_letter_start > 0:
                resume_content = result[:cover_letter_start].strip()
                cover_letter_content = result[cover_letter_start:].strip()
                
                # Clean up the headers from the content
                if cover_letter_content.lower().startswith('cover letter'):
                    cover_letter_content = '\n'.join(cover_letter_content.split('\n')[1:]).strip()
                
                print(f"Method 4 success - Resume: {len(resume_content)}, Cover Letter: {len(cover_letter_content)}")
        except Exception as e:
            print(f"Method 4 failed: {e}")
    
    print(f"Final parsing result - Resume: {len(resume_content)} chars, Cover Letter: {len(cover_letter_content)} chars")
    
    return resume_content, cover_letter_content

def optimize_resume_and_cover_letter(resume_text, job_description, user_notes):
    """Use Claude to optimize resume and generate cover letter"""
    
    prompt = f"""
You are an expert resume writer and ATS optimization specialist with deep knowledge of modern hiring practices. Your task is to transform the provided resume into a compelling, keyword-optimized document that will pass AI screening while capturing human attention.

CRITICAL RULES:
1. NEVER fabricate experiences, skills, dates, or qualifications
2. Only reorganize, enhance language, and strategically highlight existing information
3. Use user notes to add context or clarify ambiguous information
4. Ensure ATS compatibility while maintaining visual appeal
5. Focus on quantifiable achievements and impact-driven language

OPTIMIZATION STRATEGY:
- Extract and prominently feature relevant keywords from the job description
- Transform passive descriptions into active, achievement-focused statements
- Quantify accomplishments wherever possible (percentages, numbers, scale)
- Use power verbs and industry-specific terminology
- Structure content for maximum impact and readability
- Ensure keyword density matches job requirements without keyword stuffing

ORIGINAL RESUME:
{resume_text}

TARGET JOB DESCRIPTION:
{job_description}

ADDITIONAL CONTEXT:
{user_notes if user_notes else "No additional context provided"}

RESUME OPTIMIZATION GUIDELINES:

**STRUCTURE & FORMATTING:**
- Lead with a compelling Professional Summary (3-4 lines) that mirrors job requirements
- Use clear section headers: Professional Summary, Core Competencies, Professional Experience, Education, Certifications
- Prioritize most relevant sections based on job requirements
- Use consistent formatting and strategic white space

**CONTENT ENHANCEMENT:**
- Start bullet points with powerful action verbs (Spearheaded, Orchestrated, Optimized, etc.)
- Include specific metrics and outcomes where they exist in original content
- Use the STAR method framework (Situation, Task, Action, Result) for major achievements
- Incorporate relevant keywords naturally throughout all sections
- Highlight technical skills, certifications, and tools mentioned in job posting

**LANGUAGE OPTIMIZATION:**
- Replace weak phrases with strong, specific language
- Use industry terminology that matches the job description
- Ensure each bullet point demonstrates value and impact
- Vary sentence structure while maintaining clarity
- Eliminate redundancy and filler words

**ATS OPTIMIZATION:**
- Include exact keyword matches from job description
- Use standard section headings
- Avoid graphics, tables, or complex formatting
- Include both acronyms and full terms (e.g., "Search Engine Optimization (SEO)")
- Ensure critical skills appear multiple times in different contexts

Please provide your response in exactly this format:

## OPTIMIZED RESUME

[Create a strategically organized, keyword-optimized resume that transforms the original content into a compelling professional narrative. Focus on impact, achievements, and alignment with the target role while maintaining complete accuracy to the original information.]

## COVER LETTER

[Cover letter content here - 3-4 paragraphs, professional but personable]

Make sure to use exactly "## OPTIMIZED RESUME" and "## COVER LETTER" as section headers.
"""

    try:
        client = get_anthropic_client()
        
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            temperature=0.3,
            messages=[{"role": "user", "content": prompt}]
        )
        
        response_text = message.content[0].text if message.content else "Error generating response"
        print(f"Claude API response received: {len(response_text)} characters")
        
        return response_text
        
    except Exception as e:
        print(f"Anthropic API Error: {str(e)}")
        traceback.print_exc()
        return f"Error communicating with AI: {str(e)}"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/optimize', methods=['POST'])
def optimize():
    try:
        print("=== OPTIMIZE ENDPOINT CALLED ===")
        
        # Check if file was uploaded
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file uploaded'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Please upload PDF, DOC, DOCX, or TXT files.'}), 400
        
        # Get job description and notes
        job_description = request.form.get('job_description', '').strip()
        user_notes = request.form.get('user_notes', '').strip()
        
        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400
        
        print(f"File: {file.filename}, Job desc length: {len(job_description)}, Notes length: {len(user_notes)}")
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        print(f"File saved to: {file_path}")
        
        # Extract text from resume
        print("Extracting text from file...")
        resume_text = extract_text_from_file(file_path, filename)
        
        # Clean up uploaded file
        try:
            os.remove(file_path)
            print("Uploaded file cleaned up")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up file {file_path}: {cleanup_error}")
        
        if resume_text.startswith("Error"):
            print(f"Text extraction error: {resume_text}")
            return jsonify({'error': resume_text}), 400
        
        print(f"Resume text extracted: {len(resume_text)} characters")
        
        # Check API key before making request
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            print("ERROR: ANTHROPIC_API_KEY not found in environment")
            return jsonify({'error': 'AI service configuration error. Please contact support.'}), 500
        
        print("API key found, generating optimized content...")
        
        # Generate optimized content
        result = optimize_resume_and_cover_letter(resume_text, job_description, user_notes)
        
        if result.startswith("Error"):
            print(f"AI generation error: {result}")
            return jsonify({'error': result}), 500
        
        print("=== AI RESPONSE RECEIVED ===")
        print(f"Response length: {len(result)} characters")
        print(f"Response preview: {result[:300]}...")
        
        # Parse the result using enhanced parsing
        print("Parsing AI response...")
        resume_content, cover_letter_content = parse_ai_response(result)
        
        # Validate parsed content
        if not resume_content or len(resume_content.strip()) < 50:
            print(f"ERROR: Resume content is too short or empty. Length: {len(resume_content) if resume_content else 0}")
            print(f"Raw result preview: {result[:500]}...")
            return jsonify({'error': 'Could not extract valid resume content from AI response. Please try again.'}), 500
        
        if not cover_letter_content or len(cover_letter_content.strip()) < 50:
            print(f"ERROR: Cover letter content is too short or empty. Length: {len(cover_letter_content) if cover_letter_content else 0}")
            print(f"Raw result preview: {result[:500]}...")
            return jsonify({'error': 'Could not extract valid cover letter content from AI response. Please try again.'}), 500
        
        print(f"Content validation passed - Resume: {len(resume_content)} chars, Cover Letter: {len(cover_letter_content)} chars")
        
        # Make session permanent
        session.permanent = True
        
        # Store in session with verification
        print("Storing content in session...")
        session['resume_content'] = resume_content
        session['cover_letter_content'] = cover_letter_content
        
        # Verify session storage worked
        stored_resume = session.get('resume_content')
        stored_cover_letter = session.get('cover_letter_content')
        
        if not stored_resume:
            print("ERROR: Failed to store resume content in session")
            return jsonify({'error': 'Failed to store resume content. Please try again.'}), 500
        
        if not stored_cover_letter:
            print("ERROR: Failed to store cover letter content in session")
            return jsonify({'error': 'Failed to store cover letter content. Please try again.'}), 500
        
        print(f"Session storage verified - Resume: {len(stored_resume)} chars, Cover Letter: {len(stored_cover_letter)} chars")
        print("=== OPTIMIZATION COMPLETED SUCCESSFULLY ===")
        
        return jsonify({
            'success': True,
            'result': result,
            'resume_content': resume_content,
            'cover_letter_content': cover_letter_content
        })
        
    except Exception as e:
        print(f"OPTIMIZE ERROR: {str(e)}")
        print(f"Error type: {type(e)}")
        traceback.print_exc()
        
        # Return a proper JSON error response
        error_message = f'Server error: {str(e)}'
        return jsonify({'error': error_message}), 500

@app.route('/download/<file_type>')
def download_file(file_type):
    try:
        print(f"=== DOWNLOAD REQUEST: {file_type} ===")
        
        if file_type not in ['resume', 'cover_letter']:
            print(f"ERROR: Invalid file type: {file_type}")
            return jsonify({'error': 'Invalid file type'}), 400
        
        # Debug session contents
        print(f"Session keys: {list(session.keys())}")
        print(f"Session ID: {session.get('_id', 'No ID')}")
        
        # Get content from session
        content_key = f'{file_type}_content'
        content = session.get(content_key)
        
        print(f"Content for {file_type}: {'Found' if content else 'Not found'}")
        
        if not content:
            print(f"ERROR: No content found for {file_type}")
            available_keys = [k for k in session.keys() if not k.startswith('_')]
            print(f"Available session keys: {available_keys}")
            return jsonify({'error': f'No {file_type} content available. Please generate content first.'}), 400
        
        if len(content.strip()) < 10:
            print(f"ERROR: Content too short for {file_type}: {len(content)} characters")
            return jsonify({'error': f'{file_type} content is too short. Please regenerate.'}), 400
        
        print(f"Creating {file_type} document with {len(content)} characters")
        
        # Create document
        try:
            if file_type == 'resume':
                filename = create_docx_resume(content)
            else:
                filename = create_docx_cover_letter(content)
            
            filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
            
            # Verify file exists and has content
            if not os.path.exists(filepath):
                raise Exception(f"Document file was not created at {filepath}")
            
            file_size = os.path.getsize(filepath)
            if file_size == 0:
                raise Exception(f"Document file is empty: {filepath}")
            
            print(f"Document ready for download: {filename} ({file_size} bytes)")
            print(f"=== DOWNLOAD SUCCESSFUL: {file_type} ===")
            
            return send_file(
                filepath,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        except Exception as doc_error:
            print(f"Document creation error: {str(doc_error)}")
            traceback.print_exc()
            return jsonify({'error': f'Failed to create {file_type} document: {str(doc_error)}'}), 500
        
    except Exception as e:
        print(f"DOWNLOAD ERROR: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': f'Download failed: {str(e)}'}), 500

# Add error handler for 500 errors
@app.errorhandler(500)
def internal_error(error):
    print(f"500 Error Handler Called: {error}")
    return jsonify({'error': 'Internal server error occurred'}), 500

# Add error handler for 404 errors  
@app.errorhandler(404)
def not_found_error(error):
    print(f"404 Error Handler Called: {error}")
    return jsonify({'error': 'Endpoint not found'}), 404

# Add error handler for general exceptions
@app.errorhandler(Exception)
def handle_exception(e):
    print(f"General Exception Handler Called: {e}")
    print(f"Exception type: {type(e)}")
    traceback.print_exc()
    return jsonify({'error': f'Unexpected error: {str(e)}'}), 500

# Add startup logging for Render
@app.before_first_request
def startup_check():
    print("=== RENDER STARTUP CHECK ===")
    print(f"Python version: {os.sys.version}")
    print(f"Current working directory: {os.getcwd()}")
    print(f"Upload folder exists: {os.path.exists(app.config['UPLOAD_FOLDER'])}")
    print(f"Download folder exists: {os.path.exists(app.config['DOWNLOAD_FOLDER'])}")
    
    # Check for required environment variables
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    secret_key = os.environ.get("SECRET_KEY")
    
    print(f"ANTHROPIC_API_KEY set: {'Yes' if api_key else 'No'}")
    print(f"SECRET_KEY set: {'Yes' if secret_key else 'No'}")
    
    if api_key:
        print(f"API key preview: {api_key[:10]}...{api_key[-4:] if len(api_key) > 14 else '***'}")
    
    # Test Anthropic client initialization
    try:
        test_client = get_anthropic_client()
        print("✅ Anthropic client initialized successfully")
    except Exception as e:
        print(f"❌ Anthropic client initialization failed: {e}")
    
    print("=== STARTUP CHECK COMPLETE ===")

# Health check endpoint for Render with detailed info
@app.route('/health')
def health_check():
    try:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        health_data = {
            'status': 'healthy', 
            'service': 'resume-optimizer',
            'api_key_configured': bool(api_key),
            'upload_folder_exists': os.path.exists(app.config['UPLOAD_FOLDER']),
            'download_folder_exists': os.path.exists(app.config['DOWNLOAD_FOLDER'])
        }
        
        # Test Anthropic client
        try:
            get_anthropic_client()
            health_data['anthropic_client'] = 'initialized'
        except Exception as e:
            health_data['anthropic_client'] = f'error: {str(e)}'
            health_data['status'] = 'degraded'
        
        return jsonify(health_data), 200
    except Exception as e:
        return jsonify({
            'status': 'unhealthy', 
            'error': str(e)
        }), 500

if __name__ == '__main__':
    # Development server
    port = int(os.environ.get('PORT', 5000))
    print(f"Starting server on port {port}")
    run_startup_check()  # Run startup check in development
    app.run(host='0.0.0.0', port=port, debug=False)
else:
    # Production server (Gunicorn on Render)
    print("=== STARTING RESUME OPTIMIZER IN PRODUCTION MODE ===")
    print(f"Upload folder: {app.config['UPLOAD_FOLDER']}")
    print(f"Download folder: {app.config['DOWNLOAD_FOLDER']}")
    
    # Run startup check in production
    run_startup_check()
    
    print("=== PRODUCTION STARTUP COMPLETE ===")